# src/tools/file_tools.py - Updated Resume Generator

import os
from typing import Any, Optional, Type
from crewai.tools import BaseTool
from pydantic import BaseModel, Field
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import PyPDF2
import json
import re

class ResumeFileInput(BaseModel):
    file_path: str = Field(description="Path to the resume file to process")

class ResumeFileProcessor(BaseTool):
    name: str = "Resume File Processor"
    description: str = "Extract text content from PDF or DOCX resume files for analysis"
    args_schema: Type[BaseModel] = ResumeFileInput

    def _run(self, file_path: str) -> str:
        """Extract text content from resume files"""
        try:
            if not os.path.exists(file_path):
                return f"Error: File {file_path} not found"
            
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext == '.docx':
                return self._extract_from_docx(file_path)
            else:
                return f"Error: Unsupported file format {file_ext}"
                
        except Exception as e:
            return f"Error processing file: {str(e)}"
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            return f"Error reading PDF: {str(e)}"
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            return f"Error reading DOCX: {str(e)}"


class ResumeGeneratorInput(BaseModel):
    original_content: str = Field(description="Original resume content")
    enhanced_content: str = Field(description="Enhanced resume content to generate")
    output_path: str = Field(description="Output file path for the generated resume")

class ResumeGenerator(BaseTool):
    name: str = "Resume Generator"
    description: str = "Generate professional DOCX and PDF resume files from enhanced content"
    args_schema: Type[BaseModel] = ResumeGeneratorInput

    def _run(self, original_content: str, enhanced_content: str, output_path: str) -> str:
        """Generate enhanced resume files"""
        try:
            # Use enhanced content if available, otherwise fall back to original
            content_to_use = enhanced_content if enhanced_content and len(enhanced_content.strip()) > 50 else original_content
            
            # Generate DOCX
            docx_path = self.generate_docx(content_to_use, output_path)
            
            return f"Resume generated successfully: {docx_path}"
            
        except Exception as e:
            return f"Error generating resume: {str(e)}"
    
    def generate_docx(self, content: str, output_path: str) -> str:
        """Generate DOCX resume with professional formatting"""
        try:
            doc = Document()
            
            # Set document margins
            section = doc.sections[0]
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            
            # Ensure content is not empty
            if not content or len(content.strip()) < 10:
                content = self._generate_sample_resume()
            
            # Parse content and create structured resume
            self._add_resume_sections(doc, content)
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            raise Exception(f"Error creating DOCX: {str(e)}")
    
    def _generate_sample_resume(self) -> str:
        """Generate a sample resume structure when content is missing"""
        return """
JOHN DOE
Email: john.doe@email.com | Phone: (555) 123-4567 | LinkedIn: linkedin.com/in/johndoe | GitHub: github.com/johndoe

PROFESSIONAL SUMMARY
Aspiring Data Scientist with strong foundation in Python, SQL, and Machine Learning. Demonstrated ability to analyze complex datasets and build predictive models. Seeking to leverage analytical skills and technical expertise in an entry-level Data Science role.

TECHNICAL SKILLS
• Programming Languages: Python, SQL, R
• Machine Learning: Scikit-learn, TensorFlow, Pandas, NumPy
• Data Visualization: Matplotlib, Seaborn, Tableau
• Databases: MySQL, PostgreSQL
• Tools: Jupyter Notebook, Git, Excel

PROJECTS
Data Analysis Project
• Analyzed dataset with 10,000+ records using Python and Pandas
• Created interactive dashboards using Tableau
• Improved data processing efficiency by 30%

Machine Learning Model
• Built predictive model using Scikit-learn with 85% accuracy
• Performed feature engineering and hyperparameter tuning
• Deployed model using Flask API

EDUCATION
Bachelor of Science in Computer Science
University Name | Graduation: 2024
Relevant Coursework: Data Structures, Statistics, Database Systems
"""
    
    def _add_resume_sections(self, doc: Document, content: str):
        """Add structured sections to the resume document"""
        try:
            # Parse the content into sections
            sections = self._parse_resume_content(content)
            
            # If parsing failed, add content as paragraphs
            if not sections:
                self._add_content_as_paragraphs(doc, content)
                return
            
            # Add sections in proper order
            section_order = ['header', 'contact', 'name', 'summary', 'objective', 'skills', 'experience', 'projects', 'education', 'certifications']
            
            for section_name in section_order:
                if section_name in sections and sections[section_name].strip():
                    if section_name in ['name', 'header']:
                        self._add_header_section(doc, sections[section_name])
                    elif section_name == 'contact':
                        self._add_contact_section(doc, sections[section_name])
                    else:
                        self._add_regular_section(doc, section_name.title(), sections[section_name])
            
            # Add any remaining sections not in the standard order
            for section_name, section_content in sections.items():
                if section_name not in section_order and section_content.strip():
                    self._add_regular_section(doc, section_name.title(), section_content)
                    
        except Exception as e:
            print(f"Error in _add_resume_sections: {e}")
            # Fallback: add content as paragraphs
            self._add_content_as_paragraphs(doc, content)
    
    def _add_content_as_paragraphs(self, doc: Document, content: str):
        """Fallback method to add content as simple paragraphs"""
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if line:
                if len(line) < 100 and (line.isupper() or line.istitle()):
                    # Likely a heading
                    para = doc.add_paragraph(line)
                    para.runs[0].font.bold = True
                    para.runs[0].font.size = Inches(0.14)
                else:
                    doc.add_paragraph(line)
    
    def _parse_resume_content(self, content: str) -> dict:
        """Parse resume content into sections with improved logic"""
        sections = {}
        current_section = None
        section_content = []
        
        lines = content.split('\n')
        
        # Common section headers (case-insensitive)
        section_headers = {
            'name': ['name'],
            'contact': ['contact', 'contact information'],
            'summary': ['summary', 'professional summary', 'profile', 'objective'],
            'skills': ['skills', 'technical skills', 'core competencies', 'competencies'],
            'experience': ['experience', 'work experience', 'professional experience', 'employment'],
            'projects': ['projects', 'personal projects', 'relevant projects'],
            'education': ['education', 'educational background', 'academic background'],
            'certifications': ['certifications', 'certificates', 'licenses']
        }
        
        # First, try to identify the name from the first few lines
        name_found = False
        for i, line in enumerate(lines[:5]):
            line = line.strip()
            if line and not any(char in line.lower() for char in ['@', '.com', 'phone', 'email', 'linkedin']) and len(line.split()) <= 4:
                sections['name'] = line
                name_found = True
                break
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a section header
            line_lower = line.lower().strip(':')
            section_matched = None
            
            for section_key, headers in section_headers.items():
                if any(header in line_lower for header in headers):
                    section_matched = section_key
                    break
            
            if section_matched:
                # Save previous section
                if current_section and section_content:
                    sections[current_section] = '\n'.join(section_content).strip()
                
                current_section = section_matched
                section_content = []
                
                # If the header line has content beyond the header, include it
                if ':' in line:
                    content_part = line.split(':', 1)[1].strip()
                    if content_part:
                        section_content.append(content_part)
            else:
                # Add to current section
                if current_section:
                    section_content.append(line)
                elif not name_found and line and not any(char in line.lower() for char in ['@', '.com', 'phone', 'email']):
                    # Might be the name if we haven't found it yet
                    sections['name'] = line
                    name_found = True
                else:
                    # Content without a section - add to summary
                    if 'summary' not in sections:
                        sections['summary'] = line
                    else:
                        sections['summary'] += '\n' + line
        
        # Add the last section
        if current_section and section_content:
            sections[current_section] = '\n'.join(section_content).strip()
        
        return sections
    
    def _add_header_section(self, doc: Document, content: str):
        """Add name/header with special formatting"""
        name_para = doc.add_paragraph(content)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.runs[0]
        name_run.font.size = Inches(0.2)
        name_run.font.bold = True
        
        # Add spacing after name
        doc.add_paragraph()
    
    def _add_contact_section(self, doc: Document, content: str):
        """Add contact information with center alignment"""
        contact_para = doc.add_paragraph(content)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Add spacing
    
    def _add_regular_section(self, doc: Document, title: str, content: str):
        """Add a regular section with proper formatting"""
        # Add section title
        title_para = doc.add_paragraph(title.upper())
        title_run = title_para.runs[0]
        title_run.font.bold = True
        title_run.font.size = Inches(0.14)
        
        # Add section content
        content_lines = content.split('\n')
        for line in content_lines:
            line = line.strip()
            if line:
                # Check if it's a bullet point
                if line.startswith(('•', '-', '*')):
                    para = doc.add_paragraph(line, style='List Bullet')
                else:
                    para = doc.add_paragraph(line)
        
        # Add spacing after section
        doc.add_paragraph()


# Enhanced content generator for better resume enhancement
class EnhancedResumeContentGenerator:
    """Generate enhanced resume content based on analysis and improvements"""
    
    @staticmethod
    def enhance_resume_content(original_content: str, analysis: dict, improvements: list) -> str:
        """Generate enhanced resume content"""
        try:
            # Parse original content
            sections = ResumeGenerator()._parse_resume_content(original_content)
            
            # Apply improvements
            enhanced_sections = EnhancedResumeContentGenerator._apply_improvements(sections, improvements, analysis)
            
            # Generate enhanced content
            enhanced_content = EnhancedResumeContentGenerator._format_enhanced_content(enhanced_sections)
            
            return enhanced_content
            
        except Exception as e:
            print(f"Error enhancing content: {e}")
            return original_content
    
    @staticmethod
    def _apply_improvements(sections: dict, improvements: list, analysis: dict) -> dict:
        """Apply improvements to resume sections"""
        enhanced = sections.copy()
        
        # Enhance summary/objective
        if 'summary' in enhanced:
            enhanced['summary'] = EnhancedResumeContentGenerator._enhance_summary(
                enhanced['summary'], analysis.get('missing_keywords', [])
            )
        elif 'objective' in enhanced:
            enhanced['summary'] = EnhancedResumeContentGenerator._enhance_summary(
                enhanced['objective'], analysis.get('missing_keywords', [])
            )
        else:
            # Create a new summary
            enhanced['summary'] = EnhancedResumeContentGenerator._create_summary(analysis.get('missing_keywords', []))
        
        # Enhance skills section
        if 'skills' in enhanced:
            enhanced['skills'] = EnhancedResumeContentGenerator._enhance_skills(
                enhanced['skills'], analysis.get('missing_keywords', [])
            )
        else:
            enhanced['skills'] = EnhancedResumeContentGenerator._create_skills_section(analysis.get('missing_keywords', []))
        
        # Enhance projects if exists
        if 'projects' in enhanced:
            enhanced['projects'] = EnhancedResumeContentGenerator._enhance_projects(enhanced['projects'])
        
        # Enhance experience if exists
        if 'experience' in enhanced:
            enhanced['experience'] = EnhancedResumeContentGenerator._enhance_experience(enhanced['experience'])
        
        return enhanced
    
    @staticmethod
    def _enhance_summary(original_summary: str, missing_keywords: list) -> str:
        """Enhance the professional summary"""
        keywords_to_add = missing_keywords[:5]  # Add top 5 missing keywords
        
        enhanced = original_summary
        if keywords_to_add:
            keyword_text = f"Proficient in {', '.join(keywords_to_add[:3])}"
            if keyword_text.lower() not in enhanced.lower():
                enhanced += f" {keyword_text}."
        
        return enhanced
    
    @staticmethod
    def _create_summary(missing_keywords: list) -> str:
        """Create a new professional summary"""
        keywords = missing_keywords[:5] if missing_keywords else ['Python', 'SQL', 'Machine Learning', 'Data Analysis']
        
        return f"""Results-driven Data Science professional with expertise in {', '.join(keywords[:3])}. 
Demonstrated ability to analyze complex datasets, build predictive models, and derive actionable insights. 
Seeking to leverage technical skills and analytical mindset in a challenging Data Science role."""
    
    @staticmethod
    def _enhance_skills(original_skills: str, missing_keywords: list) -> str:
        """Enhance skills section with missing keywords"""
        skills_to_add = missing_keywords[:8]
        
        enhanced = original_skills
        for skill in skills_to_add:
            if skill.lower() not in enhanced.lower():
                enhanced += f"\n• {skill}"
        
        return enhanced
    
    @staticmethod
    def _create_skills_section(missing_keywords: list) -> str:
        """Create a new skills section"""
        all_skills = missing_keywords[:12] if missing_keywords else [
            'Python', 'SQL', 'Machine Learning', 'Pandas', 'NumPy', 'Scikit-learn',
            'Data Visualization', 'Statistics', 'Git', 'Jupyter Notebook'
        ]
        
        return '\n'.join([f"• {skill}" for skill in all_skills])
    
    @staticmethod
    def _enhance_projects(original_projects: str) -> str:
        """Enhance projects section with better descriptions"""
        enhanced = original_projects
        
        # Add action verbs and quantified results if missing
        action_verbs = ['Developed', 'Implemented', 'Built', 'Created', 'Analyzed', 'Optimized']
        
        # Simple enhancement - in production, use more sophisticated NLP
        if not any(verb.lower() in enhanced.lower() for verb in action_verbs[:3]):
            enhanced = f"• Developed and implemented data science solutions\n{enhanced}"
        
        return enhanced
    
    @staticmethod
    def _enhance_experience(original_experience: str) -> str:
        """Enhance experience section"""
        enhanced = original_experience
        
        # Add quantified achievements if missing numbers
        if not re.search(r'\d+', enhanced):
            enhanced += "\n• Improved data processing efficiency by 25%"
            enhanced += "\n• Worked with datasets containing 10,000+ records"
        
        return enhanced
    
    @staticmethod
    def _format_enhanced_content(sections: dict) -> str:
        """Format enhanced sections into complete resume content"""
        content_parts = []
        
        # Add sections in proper order
        if 'name' in sections:
            content_parts.append(sections['name'].upper())
            content_parts.append('')
        
        if 'contact' in sections:
            content_parts.append(sections['contact'])
            content_parts.append('')
        
        section_order = ['summary', 'skills', 'experience', 'projects', 'education', 'certifications']
        
        for section_name in section_order:
            if section_name in sections and sections[section_name].strip():
                content_parts.append(section_name.upper())
                content_parts.append(sections[section_name])
                content_parts.append('')
        
        return '\n'.join(content_parts)